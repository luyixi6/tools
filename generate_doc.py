# -*- coding: utf-8 -*-
"""生成《c++检查》项目说明文档（Word 格式）"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "c++检查.docx"

doc = Document()

# ---------- 全局样式 ----------
normal = doc.styles["Normal"]
normal.font.name = "微软雅黑"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for level, size in [(1, 16), (2, 13), (3, 11.5)]:
    style = doc.styles[f"Heading {level}"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_cn(run, font="微软雅黑"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def para(text, size=10.5, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_cn(r)
    r.font.size = Pt(10.5)
    return p


def code_block(code, lang="python"):
    """代码块：等宽字体 + 浅灰底纹"""
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        # 底纹
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F2F2F2")
        p.paragraph_format.element.get_or_add_pPr().append(shd)


# ================= 封面 =================
doc.add_paragraph()
title = para("C++ 静态检查与动态检查项目", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("（基于大语言模型的智能代码审查工具）", size=13, color=(0x55, 0x55, 0x55), align=WD_ALIGN_PARAGRAPH.CENTER)
para("CPP Inspector 项目文档", size=11, color=(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()

# ================= 一、概念 =================
doc.add_heading("一、静态检查与动态检查的概念", level=1)

doc.add_heading("1.1 静态检查（Static Analysis）", level=2)
para("静态检查是指在不运行程序的前提下，通过对源代码文本、语法树或中间表示进行解析与分析，"
     "发现代码中潜在缺陷和违规模式的技术。它关注的是\"代码本身写得对不对\"。")
para("常见的静态检查内容包括：", bold=True)
bullet("内存安全：裸指针的 new/delete 是否配对、是否存在内存泄漏、悬空指针、双重释放")
bullet("类型安全：隐式窄化转换、C 风格强制转换、有符号/无符号混用、reinterpret_cast 滥用")
bullet("异常安全：缺少 RAII、析构函数未声明 noexcept、异常路径上的资源泄漏")
bullet("const 正确性：本应为 const 的参数或成员函数")
bullet("未定义行为：未初始化变量、越界访问、空指针解引用、移动后使用")
bullet("现代 C++ 规范：缺少 override、可用算法替代的裸循环、C 风格数组等")
para("传统静态分析工具如 clang-tidy、cppcheck 依赖规则库做模式匹配；本项目则通过大语言模型"
     "理解代码语义，能够发现更深层、依赖上下文的逻辑缺陷。")

doc.add_heading("1.2 动态检查（Dynamic Analysis）", level=2)
para("动态检查是指程序在运行时（或模拟运行时）对执行路径进行检查，关注\"代码运行时会不会出错\"。"
     "本项目由于无法真正执行被测代码，采用\"模拟执行路径推演\"的方式：让大语言模型在思维中"
     "沿着代码的各个分支逐步推演，找出运行时才可能暴露的问题。")
para("本项目\"动态检查\"（模拟执行路径）关注的内容：", bold=True)
bullet("控制流：不可达代码、缺失的 switch 分支、各分支是否可达、死循环")
bullet("函数契约：前置/后置条件、参数校验、返回值处理")
bullet("并发：潜在数据竞争、缺少互斥锁、原子操作正确性")
bullet("资源生命周期：文件句柄未关闭、RAII 违规、双重关闭")
bullet("错误处理：缺少错误检查、异常被吞、错误码传播")

doc.add_heading("1.3 两者的区别", level=2)
para("静态检查侧重于代码的\"结构正确性\"，动态检查侧重于\"运行时的行为正确性\"。本项目将两者统一："
     "静态检查保证代码不违反语言规则与编码规范，动态检查保证代码在各种执行路径下不出现运行时错误。"
     "两者共同构成完整的代码质量保障。")

# ================= 二、项目概述 =================
doc.add_heading("二、项目概述", level=1)
doc.add_heading("2.1 技术架构", level=2)
para("项目采用前后端分离架构：前端 React + TypeScript 提供可视化交互界面，后端 FastAPI 提供 REST 接口"
     "与 WebSocket 实时推送，底层通过大语言模型（LLM）API 完成代码审查。")
code_block(
"""┌──────────────────────────────────────┐
│           Web 前端 (React)            │
│   项目配置 → 扫描 → 分析 → 结果审查     │
└──────────────────┬───────────────────┘
                   │ HTTP / WebSocket
┌──────────────────┴───────────────────┐
│         Python 后端 (FastAPI)         │
│  ┌────────┐ ┌────────┐ ┌───────────┐ │
│  │ 扫描器 │→│ 分析器 │→│ LLM 客户端 │ │
│  └────────┘ └────────┘ └───────────┘ │
│                 │                     │
│  ┌──────────────┴──────┐             │
│  │  修复应用 + SQLite 存储│             │
│  └─────────────────────┘             │
└──────────────────┬───────────────────┘
                   │ API 调用
┌──────────────────┴───────────────────┐
│   大语言模型 (Claude/DeepSeek/OpenAI) │
└──────────────────────────────────────┘"""
)

doc.add_heading("2.2 整体流程", level=2)
para("系统将代码审查拆解为以下 6 个阶段，形成完整闭环：", bold=True)
bullet("1. 项目扫描：递归扫描目录，收集所有 .cpp/.h 等源文件")
bullet("2. 模块分区：构建 #include 依赖图，用强连通分量（SCC）算法把项目切分为高内聚模块")
bullet("3. 逐模块分析：将每个模块的代码连同上游接口摘要发给 LLM，进行静态+动态检查")
bullet("4. 结果聚合：合并所有模块发现的问题，按严重程度/类别统计，去重")
bullet("5. 交互式审查：前端展示问题与 diff，用户逐条接受或拒绝")
bullet("6. 应用修复：接受的问题自动写回源文件（带备份，可回滚）")

# ================= 三、核心逻辑 =================
doc.add_heading("三、核心逻辑", level=1)

doc.add_heading("3.1 模块分块策略（解决大项目超上下文）", level=2)
para("大型 C++ 项目代码量往往远超 LLM 的上下文窗口，无法一次性全部发送。项目通过 include 依赖图"
     "将项目切分为多个模块，按模块逐个分析。")
para("分区采用三级自动降级策略：", bold=True)
bullet("优先：解析 #include 关系构建依赖图，用 Tarjan 算法求强连通分量（SCC）")
bullet("每个 SCC（或合并后的 SCC 组）作为一个模块，保证模块内高内聚")
bullet("每个模块获得唯一的 UUID 标识，确保多次扫描互不覆盖")
bullet("单文件过大时，按函数边界继续切分，先发接口后发实现")

doc.add_heading("3.2 上下文窗口管理", level=2)
para("每个模块的分析请求包含以下内容，以控制 token 用量：", bold=True)
bullet("依赖接口摘要：仅发送上游模块的头文件声明部分，不发送实现")
bullet("当前模块头文件：完整发送")
bullet("当前模块实现文件：分批发送，超长文件按函数切分")
bullet("传统工具诊断结果：clang-tidy / cppcheck 输出作为辅助输入")

# ================= 四、核心代码 =================
doc.add_heading("四、核心代码", level=1)

doc.add_heading("4.1 文件扫描器（scanner/file_scanner.py）", level=2)
para("递归遍历项目目录，过滤排除目录和生成文件，收集 C++ 源文件。")
code_block(
"""CPP_EXTENSIONS = {".cpp", ".cc", ".cxx", ".c++", ".C", ".c"}
HEADER_EXTENSIONS = {".h", ".hpp", ".hxx", ".h++", ".H"}

class FileScanner:
    def scan(self):
        files = []
        for dirpath, dirnames, filenames in os.walk(self.project_root):
            # 过滤排除目录（build/third_party 等）
            dirnames[:] = [d for d in dirnames
                           if d not in self.exclude_dirs]
            for fname in filenames:
                _, ext = os.path.splitext(fname)
                if ext not in ALL_SOURCE_EXTS:
                    continue
                # 记录绝对路径、相对路径、行数、大小
                files.append(FileInfo(...))
        return files"""
)

doc.add_heading("4.2 include 图与 SCC 分区（scanner/include_graph.py）", level=2)
para("解析 #include 指令构建依赖图，用 Tarjan 算法求强连通分量，再聚合成模块。每个模块用 UUID 唯一标识。")
code_block(
"""INCLUDE_RE = re.compile(r'^\\s*#\\s*include\\s*[<"]([^>"]+)[>"]')

class IncludeGraph:
    def build(self, files, project_root, include_paths):
        # 解析每个文件的 #include，解析到真实文件路径
        # 建立 nodes（依赖）与 rev_graph（被依赖）关系
        ...

    def find_sccs(self):
        # Tarjan 算法：找出强连通分量（循环依赖组）
        ...

    def group_sccs_into_modules(self, sccs, file_paths):
        # 给每个 SCC 分配 UUID，聚合成模块
        scc_id = {i: uuid.uuid4().hex[:12] for i in range(len(sccs))}
        # 相邻 SCC 合并（限制模块大小），生成模块列表
        ...
        return modules"""
)

doc.add_heading("4.3 Prompt 构建（llm/prompt_builder.py）", level=2)
para("系统提示词定义了静态检查和动态检查的完整规则，并要求 LLM 输出严格 JSON 格式。")
code_block(
"""CPP_ANALYSIS_SYSTEM_PROMPT = \"\"\"
You are an expert C++ code reviewer.

### Static Analysis
1. Memory Safety: raw new/delete, memory leaks, dangling pointers
2. Type Safety: narrowing conversions, C-style casts
3. Exception Safety: RAII, noexcept destructors
4. Const Correctness
5. Undefined Behavior: uninitialized vars, out-of-bounds access
6. Modern C++: override, algorithms, std::array/vector

### Dynamic Analysis (Simulated Execution Paths)
7. Control Flow: unreachable code, infinite loops
8. Function Contracts: pre/post conditions
9. Concurrency: data races, mutex locks
10. Resource Lifecycle: file handles, RAII violations
11. Error Handling: missing error checks

Output ONLY JSON:
{ "issues": [ { "severity": "...", "category": "...",
  "file": "...", "line_start": 42, "line_end": 45,
  "title": "...", "description": "...",
  "original_code": "...", "suggested_code": "...",
  "rule_reference": "..." } ] }
\"\"\" """
)

doc.add_heading("4.4 多模型客户端工厂（llm/client_factory.py）", level=2)
para("支持 Anthropic、OpenAI、DeepSeek、智谱、通义、Moonshot 等 7 种供应商，统一接口。")
code_block(
"""SUPPORTED_PROVIDERS = {
    "anthropic": {"default_model": "claude-sonnet-4-20250514"},
    "openai":    {"default_model": "gpt-4o",
                  "base_url": "https://api.openai.com/v1"},
    "deepseek":  {"default_model": "deepseek-chat",
                  "base_url": "https://api.deepseek.com/v1"},
    "zhipu":     {"default_model": "glm-4-plus", ...},
    "qwen":      {"default_model": "qwen-plus", ...},
    "moonshot":  {"default_model": "moonshot-v1-8k", ...},
}

def create_client(provider, api_key, model, max_tokens, base_url):
    # 根据 provider 返回对应的客户端实例
    # OpenAI 系走 OpenAI 兼容接口，Anthropic 走专用 SDK
    ..."""
)

doc.add_heading("4.5 模块分析调度（analyzer/static_analyzer.py）", level=2)
para("读取模块文件、估算 token、切分大文件、调用 LLM 并解析结果。")
code_block(
"""class ModuleAnalyzer:
    async def analyze_module(self, module):
        # 1. 读取模块内所有文件内容
        # 2. 估算 token，超限则按函数边界切分
        # 3. 收集上游依赖的接口摘要
        # 4. 构建 Prompt，调用 LLM
        system_prompt, user_content = build_module_prompt(
            module["name"], batch_files, dependencies,
            analysis_types, language=self.config.language)
        response = await self.client.analyze_code(...)
        return parse_analysis_response(response)"""
)

doc.add_heading("4.6 并发分析（main.py）", level=2)
para("使用 asyncio.Semaphore 限制并发数，多模块并行分析，大幅缩短总耗时。")
code_block(
"""async def _run_analysis(scan_id, module_ids):
    semaphore = asyncio.Semaphore(concurrent_modules)

    async def analyze_one(idx, module):
        async with semaphore:
            issues = await analyzer.analyze_module(module)
            insert_issues(issues)
            # 广播进度（WebSocket）
            await _broadcast(scan_id, {...})

    tasks = [asyncio.create_task(analyze_one(i, m))
             for i, m in enumerate(selected)]
    await asyncio.gather(*tasks)
    update_scan_status(scan_id, "completed")"""
)

# ================= 五、如何检查 =================
doc.add_heading("五、项目如何进行静态检查和动态检查", level=1)

doc.add_heading("5.1 检查规则的落地", level=2)
para("项目的检查规则完全体现在系统提示词（System Prompt）中。LLM 依据提示词中定义的 11 类规则"
     "（6 类静态 + 5 类动态）对每个模块进行审查。")
para("传统工具辅助：", bold=True)
bullet("clang-tidy：基于编译数据库的静态分析，结果作为 LLM 的补充输入")
bullet("cppcheck：独立静态分析工具，检测内存、性能、可移植性问题")
para("这些工具的诊断结果会附加到模块代码之后，帮助 LLM 聚焦可疑代码区域，提高诊断准确率。")

doc.add_heading("5.2 一次完整检查的执行过程", level=2)
bullet("1. 前端发起扫描请求，后端递归扫描项目文件")
bullet("2. 构建 include 图，Tarjan 算法求 SCC，切分为模块")
bullet("3. 用户在前端勾选要分析的模块，发起分析")
bullet("4. 后端并发逐模块调用 LLM，进行静态+动态检查")
bullet("5. LLM 返回 JSON 问题列表（严重程度、类别、行号、原始代码、建议代码）")
bullet("6. 结果存入 SQLite，前端展示问题列表与 diff")
bullet("7. 用户接受修复 → 代码写回源文件并备份；拒绝则跳过")

# ================= 六、如何使用 =================
doc.add_heading("六、如何使用这个项目", level=1)

doc.add_heading("6.1 环境准备", level=2)
para("依赖：Python 3.10+、Node.js 16+。", bold=True)
code_block(
"""# 安装后端依赖
pip install -r backend/requirements.txt

# 安装前端依赖并构建
cd frontend
npm install
npm run build"""
)

doc.add_heading("6.2 配置 API Key", level=2)
para("编辑项目根目录下的 config.yaml 文件，配置模型供应商、API 密钥和项目路径。")
code_block(
"""api:
  provider: deepseek            # 可选 anthropic/openai/deepseek 等
  api_key: "sk-xxxxxxxx"        # 你的 API 密钥
  model: "deepseek-chat"
  base_url: ""                  # OpenAI 兼容接口可自定义

project:
  root: "D:/your-cpp-project"   # 要检查的 C++ 项目路径

batch:
  concurrent_modules: 2         # 并发分析模块数
  rate_limit_rpm: 50            # 每分钟请求上限"""
)
para("也可以直接在 Web 界面的\"配置\"页填写 API Key 和项目路径，保存后会自动写入 config.yaml。")

doc.add_heading("6.3 启动服务", level=2)
code_block(
"""python -m uvicorn backend.main:app --host 127.0.0.1 --port 8000"""
)
para("浏览器打开 http://127.0.0.1:8000 即可使用。")

doc.add_heading("6.4 使用步骤", level=2)
bullet("1. 配置页：选择模型供应商，填写 API Key 和项目根目录，点击\"保存配置\"")
bullet("2. 扫描页：点击\"扫描项目\"，系统自动分区为多个模块")
bullet("3. 勾选需要检查的模块，点击\"分析 N 个模块\"")
bullet("4. 分析进度页实时显示进度，完成后自动进入结果页")
bullet("5. 结果页：左侧问题列表，右侧 diff 对比，点击\"接受并应用修复\"或\"拒绝\"")
bullet("6. 历史记录：顶部标签显示每次扫描的时间，可随时切换回看并继续修改")

para("")
para("说明：本项目的\"动态检查\"是通过大语言模型模拟执行路径推演实现的逻辑检查；"
     "如需真正的运行时检查，可在 config.yaml 中启用 clang-tidy/cppcheck 工具辅助，"
     "或将编译产物接入 AddressSanitizer/Valgrind 等动态分析工具。", size=9.5, color=(0x88, 0x88, 0x88))

doc.save(OUTPUT)
print(f"已生成文档：{OUTPUT}")
